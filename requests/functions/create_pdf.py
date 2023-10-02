#!/usr/bin/python3

""" Takes a docx file with placeholders and fills in placeholders and creates a PDF

"""

import os
from pathlib import Path
from docx import Document
from python_docx_replace import docx_replace, docx_get_keys
import pexpect
import re
import docx2txt
import constants as C
from fnmatch import fnmatch


# TODO: check env variables are available

from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm


class CreatePDF:
    """Extract information from docx template and also create PDF with
    required information

    """

    def __init__(self, template_dir: str | None, output_path: str | None) -> None:
        if template_dir == None:
            raise ValueError(f"No template directory stored in environmental file .env")

        if not os.path.isdir(template_dir):
            raise ValueError(f'Template directory "{ template_dir }" does not exist!')

        if output_path == None:
            raise ValueError(f"No output path stored in environmental file .env")

        # No need to check if output_path exists, as it will be created
        # in self.create()

        self.template_dir: str = str(template_dir)
        self.output_path: str = str(output_path)
        # self.template_dir: str = '/requests/templates/'
        # self.output_path: str = '/archive/'
        return None

    def check_template(
        self, search_location: str = "/requests/templates/", verbose: bool = True
    ) -> dict[str, str]:
        """Check that a single or folder worth of templates are correctly
        formatted for using placeholders

        """

        files_to_check: dict = {}
        doc: str = ""
        # TODO: doc_Regex initialise and hint type
        left_stake: list = []
        right_stake: list = []
        request: list[str] = []
        location: str = ""
        placeholders: list[list[str]] = []
        path_split: str = ""
        config_found: bool = False

        if os.path.isdir(search_location):
            for path, subdirs, files in os.walk(search_location):
                for name in files:
                    if fnmatch(name, "*.docx") and not name.startswith("~"):
                        files_to_check[os.path.join(path, name)] = ""
        elif os.path.isfile(search_location):
            files_to_check[search_location] = ""
        else:
            raise RuntimeError(
                f'Search location "{ search_location }" is'
                f"neither a valid directory or location"
            )

        for file in files_to_check:
            if len(request):
                request.pop()
            config_found = False

            doc = docx2txt.process(file)
            doc_Regex = re.compile(r"\$\{")
            left_stake = doc_Regex.findall(doc)
            doc_Regex = re.compile(r"\}")
            right_stake = doc_Regex.findall(doc)

            if len(left_stake) == len(right_stake):
                files_to_check[file] = "."
            else:
                files_to_check[file] = "F"

            path_split = file.split("/")
            request.append(path_split[-1].replace(".docx", ""))
            location = path_split[-2]
            placeholders = self.get_placeholders(location, request)

            for ph in placeholders:
                if ph[2] == "configuration":
                    config_found = True

            if config_found:
                files_to_check[file] += "."
            else:
                files_to_check[file] += "W"

        if verbose:
            print(files_to_check)

        return files_to_check

    def get_locations(self) -> list:
        """Get the 'locations' (physical or virtual) for a collection of
        different clinical requests

        """
        locations: list[str] = []

        # List of all content in a directory, filtered so only directories
        # are returned
        locations = [
            directory
            for directory in os.listdir(self.template_dir)
            if os.path.isdir(self.template_dir + directory)
        ]

        if not locations:
            raise RuntimeError(
                f'No locations folders found in "{ self.template_dir }" \
                template directory'
            )

        return locations

    # TODO: Perhaps can use docx_get_keys() instead to get keys
    # (maybe will use python-docs-tempate)
    def get_types(self, location: str) -> list[str]:
        """Get the types of tests available for a location"""
        requests: list[str] = []
        requests_path = f"{self.template_dir}{location}/"

        if not os.path.isdir(requests_path):
            raise RuntimeError(f'"{requests_path}" is not a valid location directory!')

        requests = [
            f
            for f in os.listdir(requests_path)
            if os.path.isfile(f"{requests_path}{f}")
            and f.endswith(".docx")
            and not f.startswith("~")
        ]

        requests = [x.removesuffix(".docx") for x in requests]

        if not requests:
            raise RuntimeError(f'No request requests found in "{ location }" folder')

        return requests

    def get_placeholders(self, location: str, requests: list[str]) -> list[list[str]]:
        """Gets placeholders and remove any duplicates as they appear

            Args:
            location: locations of requests (either physical or virtual)
            *argv: one or more request types
        returns:
            list[str]
        """

        # TODO: #4 might want to make this a dictionary of lists
        # (for easier searching)
        placeholders_final: list[list[str]] = []
        raw_placeholders: list[str] = []
        cleaned_placeholder: str = ""

        types_path = f"{self.template_dir}{location}/"
        template_path: str = ""

        if not os.path.isdir(types_path):
            raise RuntimeError(f'"{ types_path }" is not a valid location directory!')

        if not requests:
            raise RuntimeError(f"No requests have been specified")

        for request in requests:
            template_path = f"{ types_path }{ request }.docx"
            if not os.path.exists(template_path):
                raise RuntimeError(
                    f'Template docx file "{ template_path }" does not exist!'
                )

        for request in requests:
            doc = docx2txt.process(f"{ types_path }{ request }.docx")
            doc_Regex = re.compile(r"\$\{.*?\}")
            raw_placeholders = doc_Regex.findall(doc)

            for raw_placeholder in raw_placeholders:
                cleaned_placeholder = re.sub("[${}]", "", raw_placeholder)

                if not placeholders_final:  # Empty list
                    self.sub_list_extract(placeholders_final, cleaned_placeholder)
                elif not any(cleaned_placeholder in ph for ph in placeholders_final):
                    self.sub_list_extract(placeholders_final, cleaned_placeholder)

        return placeholders_final

    def sub_list_extract(
        self, placeholders_final: list[list[str]], cleaned_placeholder: str
    ) -> None:
        """sub routine to split via piped delimiter"""

        subList: list[str] = []

        subList.append(cleaned_placeholder)
        subList.append("")

        splitKey = cleaned_placeholder.split(C.DELIMITER)
        for s in splitKey:
            subList.append(s.strip())

        placeholders_final.append(list(subList))

        return None

    # TODO: could rewrite this to use the python-docx-template library
    # TODO: take a list of requests to process
    def create(
        self,
        location: str,
        request: str,
        placeholders: list[list[str]],
        demographics: str,
    ) -> str:
        """Actually create the PDF"""

        requests_path: str = f"{self.template_dir}{location}/"
        template_path: str = f"{ requests_path }{ request }.docx"
        temp_docx_dir: str = f"{ self.output_path }{ location }/temp/"
        temp_docx_path: str = f"{ temp_docx_dir }{ request }_" f"{ demographics }_"

        n: int = 1
        # Need 'pdf_dir' for libreOffice arguments.
        pdf_dir: str = ""
        pdf_path: str = ""
        placeholders_dict: dict[str, str] = {}
        placeholder_temp: str = ""

        # TODO: hint type libreoffice_output: type[pexpect]

        if not os.path.isdir(requests_path):
            raise RuntimeError(
                f'"{ requests_path }" is not a valid location directory!'
            )

        # TODO: need this to use a list
        if not request:
            # raise RuntimeError(f'No requests have been specified')
            raise RuntimeError(f"No request has been specified")

        if not os.path.exists(template_path):
            raise RuntimeError(
                f'Template docx file "{ template_path }" does not exist!'
            )

        if not placeholders:
            raise RuntimeError(f"No placeholders provided")

        if any(illegal in demographics for illegal in '<>?:"/\\|?*,'):
            raise RuntimeError(
                f'Illegal character in demographics - "{ demographics }"'
            )

        for x in range(len(placeholders)):
            if len(placeholders[x]) >= 4:
                if placeholders[x][3] == "picture":
                    if not os.path.exists(placeholders[x][1]):
                        raise RuntimeError(
                            f'Image filename "{ placeholders[x][1] }"' f"does not exist"
                        )

        # Make directory with sub-folders if needed
        if not os.path.isdir(temp_docx_dir):
            Path(temp_docx_dir).mkdir(parents=True, exist_ok=True)

        pdf_dir = f"{ self.output_path }{ location }"
        pdf_path = f"{ pdf_dir }/{ request }_{ demographics }_"

        while n < 10000:
            if os.path.exists(f"{ pdf_path }{ n }.pdf"):
                n = n + 1
            else:
                break

        if n >= 10000:
            raise RuntimeError("Filename increment over ran!")

        temp_docx_path = f"{ temp_docx_path }{ n }.docx"
        pdf_path = f"{ pdf_path }{ n }.pdf"

        for x in range(len(placeholders)):
            placeholder_temp = placeholders[x][1]

            if len(placeholders[x]) >= 4:
                if placeholders[x][3] == "picture":
                    placeholder_temp = f"{{{{ { placeholders[x][2] } }}}}"

            placeholders_dict[placeholders[x][0]] = placeholder_temp

        try:
            doc = Document(template_path)
            docx_replace(doc, **placeholders_dict)
            doc.save(temp_docx_path)
        except:
            raise Exception(f'Could not create the .docx file "{ temp_docx_path }"!')

        # Double check that the file has been created
        if not os.path.isfile(temp_docx_path):
            raise Exception(f'Error - file {temp_docx_path}" has not been created!')

        # TODO: need to have picture adding in placeholders returned
        # (placeholders)

        for x in range(len(placeholders)):
            if len(placeholders[x]) >= 4:
                if placeholders[x][3] == "picture":
                    self.add_picture(
                        temp_docx_path, placeholders[x][1], placeholders[x][2]
                    )

        libreoffice_output = pexpect.spawn(
            f'libreoffice --headless --convert-to pdf "{ temp_docx_path }" '
            f'--outdir "{ pdf_dir }"'
        )

        # TODO: print(type(libreoffice_output))

        if libreoffice_output.read()[0:7] != b"convert":
            raise Exception(f"Error with PDF creation via LibreOffice")

        os.remove(temp_docx_path)

        # TODO: will need to output a list
        return pdf_path

    # TODO: May remove this function later
    def add_picture(
        self, file: str, image: str, placeholder: str = "signature", width=Mm(20)
    ) -> None:
        """For adding pictures to docx. May be able to join this with
        'create'by using the python-docx-template module

        """

        doc: type[docxtpl]

        if not os.path.exists(file):
            raise Exception(f'Docx filename "{ file }" does not exist')

        if not os.path.exists(image):
            raise Exception(f'Image filename "{ image }" does not exist')

        doc = DocxTemplate(file)
        context = {placeholder: InlineImage(doc, image, width=width)}
        doc.render(context)
        doc.save(file)

        return None


if __name__ == "__main__":
    print("Running...")

    docxPtr = CreatePDF(C.TEMPLATE_DIR, C.PDF_DIR)
    # Get the different locations
    # print(f'Locations: { docxPtr.get_locations() }')

    # Get the types of requests in a the Salisbury trust
    # print(f'Available requests in Salisbury:
    # { docxPtr.get_types("Salisbury") }')

    # Get variables from
    # requests = ['Bronchoscopy 1', 'Bronchoscopy 2', 'Bronchoscopy 3']
    variables = docxPtr.get_placeholders("Salisbury", ["Lung function test"])
    # print(variables)

    # These might not work any more as Lists instead of Dic are bring used.

    variables[0][1] = "123456"
    variables[2][1] = "John"
    variables[3][1] = "Smith"
    variables[5][1] = C.UNCHECKED
    variables[6][1] = C.CHECKED

    """print('Placeholders for Lung function tests at Salisbury are:')
    for v in variables:
        print(v)
    """
    pdf_path = docxPtr.create(
        "Salisbury",
        "Lung function test",
        variables,
        "Smith_John_1234567",
        "/requests/signatures/Mark Bailey.jpeg",
    )
    # print(pdf_path)

    # docxPtr.add_picture('/requests/testing/Lung function test.docx',
    #                    '/requests/signatures/Mark Bailey.jpeg')
